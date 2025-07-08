import sys
import os
sys.path.append(os.path.abspath(os.path.join(os.path.dirname(__file__), '..')))
import pytest
import asyncio
import re
from io import BytesIO
import docx
import fitz  # PyMuPDF
from app.states.rag_state import RAGState
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter



@pytest.fixture
def estado():
    return RAGState()

class MockRAGState:
    def __init__(self):
        self.pregunta = ""
        self.respuesta = ""
        self.error_message = ""
        self.is_loading = False

    def generar(self):
        self.is_loading = True
        self.respuesta = ""
        self.error_message = ""

        texto = self.pregunta.strip()
        # Validación 1: pregunta vacía
        if not texto:
            mensaje = "Por favor, escribe una pregunta."
            self.respuesta = mensaje
            self.error_message = mensaje
            self.is_loading = False
            yield
            return
        
        # Validación 2: Validacion de emojis extendida
        if re.search(r"[\U0001F600-\U0001F64F"  # Emoticons
                    r"\U0001F300-\U0001F5FF"  # Símbolos y pictogramas
                    r"\U0001F680-\U0001F6FF"  # Transporte y mapas
                    r"\U0001F1E0-\U0001F1FF"  # Banderas
                    r"\U0001F900-\U0001F9FF"  # Emojis adicionales
                    r"\U0001FA70-\U0001FAFF"  # Emoji expansión 2020
                    r"\u2600-\u26FF"          # Misceláneos símbolos
                    r"\u2700-\u27BF"          # Dingbats
                    "]", texto):
            mensaje = "No se permiten emojis en la pregunta."
            self.respuesta = mensaje
            self.error_message = mensaje
            self.is_loading = False
            yield
            return        
       
        
        # Validación 3: contiene caracteres no latinos
        if re.search(r"[^\u0000-\u007FáéíóúÁÉÍÓÚñÑüÜçÇ\s.,;:?!¿¡()\"'-]", texto):
            mensaje = "Solo se permite ingresar texto en alfabeto latino."
            self.respuesta = mensaje
            self.error_message = mensaje
            self.is_loading = False
            yield
            return
        
        # Validación 4: no contiene ninguna letra
        if not re.search(r"[a-zA-ZáéíóúÁÉÍÓÚñÑüÜ]", texto):
            mensaje = "La pregunta debe contener letras."
            self.respuesta = mensaje
            self.error_message = mensaje
            self.is_loading = False
            yield
            return
        
        # Validación 5: menos de 3 palabras
        if len(texto.split()) < 3:
            mensaje = "La pregunta debe contener al menos 3 palabras."
            self.respuesta = mensaje
            self.error_message = mensaje
            self.is_loading = False
            yield
            return

        self.respuesta = "Respuesta simulada."
        self.error_message = ""
        self.is_loading = False    
        yield


def test_pregunta_vacia(estado):
    estado = MockRAGState()
    estado.pregunta = ""
    for _ in estado.generar():
        pass
    assert estado.error_message.lower().startswith("por favor")

def test_pregunta_valida_generica(estado):
    estado.pregunta = "¿Qué es Clean Language?"
    estado.generar()
    # Puede cambiar según tu lógica exacta
    assert estado.error_message == "" or "No hay suficiente información" in estado.respuesta

def test_pregunta_con_emojis():
    estado = MockRAGState()
    estado.pregunta = "¿Cuál es el resultado? 🤔"
    for _ in estado.generar():
        pass
    assert "emojis" in estado.error_message.lower()

def test_pregunta_con_chino():
    estado = MockRAGState()
    estado.pregunta = "这是中文"  # chino
    for _ in estado.generar():
        pass
    print("Mensaje recibido:", estado.error_message)
    assert "alfabeto latino" in estado.error_message.lower()      

def test_pregunta_sin_letras():
    estado = MockRAGState()
    estado.pregunta = "123456789"
    for _ in estado.generar():
        pass
    assert "contener letras" in estado.error_message.lower()

def test_pregunta_corta():
    estado = MockRAGState()
    estado.pregunta = "Hola mundo"
    for _ in estado.generar():
        pass
    assert "al menos 3 palabras" in estado.error_message.lower()    



# Simula un archivo y llama a tu lógica real de procesamiento
def procesar_archivo_simulado(nombre_archivo: str, contenido: bytes) -> str:
    if nombre_archivo.endswith(".txt"):
        return contenido.decode("utf-8")
    
    elif nombre_archivo.endswith(".docx"):
        file_like = BytesIO(contenido)
        documento = docx.Document(file_like)
        return "\n".join([p.text for p in documento.paragraphs])
    
    elif nombre_archivo.endswith(".pdf"):
        file_like = BytesIO(contenido)
        doc = fitz.open(stream=file_like.read(), filetype="pdf")
        texto = ""
        for pagina in doc:
            texto += pagina.get_text()
        return texto.strip()
    
    else:
        raise ValueError("Tipo de archivo no soportado.")    
    

def test_procesar_txt_simulado():
    contenido = "Esto es una prueba de texto plano.\nSegunda línea.".encode("utf-8")
    texto = procesar_archivo_simulado("archivo.txt", contenido)
    assert "prueba de texto" in texto

def test_procesar_docx_simulado():
    # Crear un archivo docx en memoria
    archivo_docx = BytesIO()
    documento = docx.Document()
    documento.add_paragraph("Este es un documento Word de prueba.")
    documento.save(archivo_docx)

    texto = procesar_archivo_simulado("archivo.docx", archivo_docx.getvalue())
    assert "Word de prueba" in texto


def test_procesar_pdf_simulado():
     
    # Crear un PDF en memoria
    pdf_bytes = BytesIO()
    c = canvas.Canvas(pdf_bytes, pagesize=letter)
    c.drawString(100, 750, "Este es un PDF de prueba.")
    c.save()
    pdf_bytes.seek(0)

    texto = procesar_archivo_simulado("archivo.pdf", pdf_bytes.read())
    assert "PDF de prueba" in texto


def test_procesar_txt_vacio():
    contenido = "".encode("utf-8")
    texto = procesar_archivo_simulado("archivo.txt", contenido)
    assert texto.strip() == ""


def test_procesar_extension_no_valida():
    contenido = b"contenido ficticio"
    with pytest.raises(ValueError) as e:
        procesar_archivo_simulado("archivo.exe", contenido)
    assert "no soportado" in str(e.value).lower()


def test_procesar_docx_vacio():
    archivo_docx = BytesIO()
    documento = docx.Document()
    documento.save(archivo_docx)

    texto = procesar_archivo_simulado("archivo.docx", archivo_docx.getvalue())
    assert texto.strip() == ""


def test_procesar_pdf_sin_texto():
       
    pdf_bytes = BytesIO()
    c = canvas.Canvas(pdf_bytes, pagesize=letter)
    # Dibujar solo una línea, sin texto
    c.line(100, 750, 300, 750)
    c.save()
    pdf_bytes.seek(0)

    texto = procesar_archivo_simulado("archivo.pdf", pdf_bytes.read())
    assert texto.strip() == ""


def test_txt_utf8_caracteredeactivates_latinos():
    contenido = "áéíóú ñ Ñ ç ü".encode("utf-8")
    texto = procesar_archivo_simulado("archivo.txt", contenido)
    assert "ñ" in texto

